
import streamlit as st
from docx import Document
from docx.text.paragraph import Paragraph
from io import BytesIO
from rapidfuzz import fuzz
import re, datetime, os, httpx

# --------------------------
# Config
# --------------------------
st.set_page_config(page_title="ATS Resume Optimizer", page_icon="✅", layout="wide")

# --------------------------
# Utility & Text helpers
# --------------------------
STOPWORDS = set({
    "the","and","for","with","that","this","from","your","you","are","our","their","they","them",
    "into","over","under","on","in","to","of","a","an","by","as","at","or","is","be","been","it",
    "we","will","can","able","through","across","such","within","including","include","etc","per",
    "using","use","vs","via","may","like"
})

CURATED_QA_KEYWORDS = {
    "exploratory","regression","functional","integration","ui","ux","user interface","test plan","defect",
    "bug","root cause","sql","logs","sdlc","agile","uat","release","pre-release","playwright","cypress",
    "selenium","automation","automated tests","manual testing","guidewire","quality","inventory",
    "stakeholders","product managers","developers","designers","customer-facing","architecture",
    "software architecture","quantitative","analytical","logical","methodical","skeptical"
}

SECTION_HEADS = [
    "PROFILE","SUMMARY","CORE SKILLS","TESTING & QA","AUTOMATION & TOOLS","DATA & ANALYTICS",
    "EDUCATION","EMPLOYMENT EXPERIENCE","EXPERIENCE","PERSONAL PROJECTS","PROJECTS",
    "PROFESSIONAL DEVELOPMENT","CERTIFICATIONS","ACCOMPLISHMENTS","AWARDS","INTERESTS"
]

def read_docx_text(doc: Document) -> str:
    parts = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    parts.append(p.text)
    return "\n".join(parts)

def get_heading_indices(doc: Document):
    idxs = {}
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip().upper()
        if t in SECTION_HEADS:
            idxs.setdefault(t, []).append(i)
    return idxs

def find_first_heading(doc: Document, name: str):
    idxs = get_heading_indices(doc)
    arr = idxs.get(name.upper(), [])
    return arr[0] if arr else None

def next_section_index(doc: Document, after_idx: int):
    for i in range(after_idx+1, len(doc.paragraphs)):
        if (doc.paragraphs[i].text or "").strip().upper() in SECTION_HEADS:
            return i
    return len(doc.paragraphs)

def delete_paragraph(paragraph: Paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None

def keep_first_n_bullets_in_block(doc: Document, start_idx: int, n_keep: int = 3):
    end = next_section_index(doc, start_idx)
    kept = 0
    i = start_idx + 1
    while i < end:
        p = doc.paragraphs[i]
        txt = (p.text or "").strip()
        if not txt:
            delete_paragraph(p)
            end -= 1
            continue
        if kept < n_keep:
            kept += 1
            i += 1
        else:
            delete_paragraph(p)
            end -= 1

def extract_keywords(jd_text: str, top_k: int = 44):
    text = re.sub(r"[^a-z0-9\s\-/+&]", " ", jd_text.lower())
    tokens = [t for t in text.split() if t not in STOPWORDS and len(t) > 2]
    freq = {}
    for t in tokens:
        freq[t] = freq.get(t, 0) + 1
    bigrams = [" ".join([tokens[i], tokens[i+1]]) for i in range(len(tokens)-1)]
    for b in bigrams:
        if all(w not in STOPWORDS for w in b.split()):
            freq[b] = freq.get(b, 0) + 1
    for k in CURATED_QA_KEYWORDS:
        freq[k] = freq.get(k, 0) + 2
    items = sorted(freq.items(), key=lambda x: x[1], reverse=True)
    return [w for w,_ in items[:top_k]]

def score_text(text: str, keywords):
    text_l = text.lower()
    matched = set()
    for kw in keywords:
        if kw in text_l or fuzz.partial_ratio(kw, text_l) >= 86:
            matched.add(kw)
    score = round(100 * len(matched) / max(1, len(keywords)), 2)
    return score, matched

def build_compact_summary(jd_keywords):
    parts = []
    parts.append("Results-driven QA Engineer with 8+ years across Insurance and Healthcare.")
    parts.append("Skilled in manual, automated, and exploratory testing of desktop, web, and mobile apps.")
    parts.append("Strong in defect management, SQL, logs, and debugging complex architectures; improve release readiness.")
    parts.append("Collaborate with product managers, developers, and designers to build the right thing, the right way.")
    parts.append("Disciplined, analytical, and methodical; comfortable making pre-release calls and pushing for quality.")
    parts.append("Completing a Master’s in Data Science and Analytics at the University of Calgary.")
    return " ".join(parts)

def trim_projects(doc: Document, heading="PERSONAL PROJECTS", keep=4):
    h = find_first_heading(doc, heading)
    if h is None:
        return
    end = next_section_index(doc, h)
    kept = 0
    i = h + 1
    while i < end:
        p = doc.paragraphs[i]
        if (p.text or "").strip():
            if kept < keep:
                kept += 1
                i += 1
            else:
                delete_paragraph(p)
                end -= 1
        else:
            delete_paragraph(p)
            end -= 1

def trim_awards(doc: Document, heading="ACCOMPLISHMENTS", keep=4):
    h = find_first_heading(doc, heading)
    if h is None: return
    end = next_section_index(doc, h)
    kept = 0; i = h+1
    while i < end:
        p = doc.paragraphs[i]
        if (p.text or "").strip():
            if kept < keep:
                kept += 1; i += 1
            else:
                delete_paragraph(p); end -= 1
        else:
            delete_paragraph(p); end -= 1

def merge_profdev_interests(doc: Document):
    pd = find_first_heading(doc, "PROFESSIONAL DEVELOPMENT")
    intr = find_first_heading(doc, "INTERESTS")
    if pd is None or intr is None:
        return
    doc.paragraphs[pd].text = "PROFESSIONAL DEVELOPMENT & INTERESTS"
    end_intr = next_section_index(doc, intr)
    interest_lines = []
    for i in range(intr+1, end_intr):
        txt = (doc.paragraphs[i].text or "").strip()
        if txt:
            interest_lines.append(txt)
    merged_line = " | ".join(interest_lines)[:400] if interest_lines else ""
    # Append merged line near PD end
    insert_at = next_section_index(doc, pd)
    # Add at the end of doc (simpler/safer to avoid index issues)
    if merged_line:
        doc.add_paragraph(merged_line)
    # Remove old interests section
    for i in range(intr, end_intr):
        delete_paragraph(doc.paragraphs[intr])

def cap_bullets_each_role(doc: Document, n_keep=3):
    exp = find_first_heading(doc, "EMPLOYMENT EXPERIENCE") or find_first_heading(doc, "EXPERIENCE")
    if exp is None: return
    i = exp + 1
    while i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        t = (p.text or "").strip()
        if t.upper() in SECTION_HEADS:
            break
        if t and (len(t) <= 120 and not t.endswith(".")):
            keep_first_n_bullets_in_block(doc, i, n_keep=n_keep)
        i += 1

def replace_profile_summary(doc: Document, new_summary: str):
    """
    Replaces PROFILE/SUMMARY content with `new_summary`.
    Uses insert-before-next-heading (no insert_paragraph_after).
    """
    h = find_first_heading(doc, "PROFILE") or find_first_heading(doc, "SUMMARY")
    if h is None:
        # No profile: create one at the top safely using insert_paragraph_before (this *is* supported)
        if doc.paragraphs:
            first = doc.paragraphs[0]
            first.insert_paragraph_before(new_summary)
            first.insert_paragraph_before("PROFILE")
        else:
            # extremely rare: empty doc
            doc.add_paragraph("PROFILE")
            doc.add_paragraph(new_summary)
        return

    # Find where this section ends
    end = next_section_index(doc, h)

    # Remove everything between heading and next heading
    i = h + 1
    while i < end:
        delete_paragraph(doc.paragraphs[i])
        end -= 1

    # After deletion, `end` now points to the next heading (or end of doc).
    # Insert the new summary **before** that next heading. This effectively puts the
    # summary right after the PROFILE/SUMMARY heading without using insert_paragraph_after.
    if end < len(doc.paragraphs):
        doc.paragraphs[end].insert_paragraph_before(new_summary)
    else:
        # if there is no next heading, just append
        doc.add_paragraph(new_summary)


def enforce_two_pages_soft(doc: Document, hard_char_cap=9500):
    txt = read_docx_text(doc)
    if len(txt) <= hard_char_cap:
        return
    trim_projects(doc, keep=3)
    trim_awards(doc, keep=3)
    cap_bullets_each_role(doc, n_keep=2)

# --------------------------
# OpenRouter LLM helpers
# --------------------------
def call_openrouter(messages, model="openrouter/auto", temperature=0.2, max_tokens=400, api_key=None, timeout=60):
    if not api_key:
        raise RuntimeError("OpenRouter API key not set. Add it to Streamlit Secrets as OPENROUTER_API_KEY.")
    url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "HTTP-Referer": "https://streamlit.io",
        "X-Title": "ATS Resume Optimizer",
        "Content-Type": "application/json",
    }
    payload = {
        "model": model,
        "messages": messages,
        "temperature": float(temperature),
        "max_tokens": int(max_tokens),
    }
    with httpx.Client(timeout=timeout) as client:
        r = client.post(url, headers=headers, json=payload)
        r.raise_for_status()
        data = r.json()
        return data["choices"][0]["message"]["content"]

def rewrite_summary_with_llm(summary_hint, jd_text, api_key, model, temperature=0.2):
    system = "You are a resume optimizer. Compress the profile summary into 5–6 lines, keep factual accuracy, and include QA/ATS keywords from the job description. Keep a neutral, professional tone. Do not use first person."
    user = f"Job Description:\n{jd_text}\n\nExisting Summary:\n{summary_hint}\n\nRewrite the summary to 5–6 concise lines. Output plain text only."
    return call_openrouter(
        [{"role":"system","content":system},{"role":"user","content":user}],
        model=model, temperature=temperature, max_tokens=240, api_key=api_key
    )

def polish_bullets_with_llm(role_block_text, jd_text, api_key, model, temperature=0.2):
    system = "You are a resume editor for QA roles. Rewrite bullets to be concise, impact-driven, and keyword-aligned. Preserve factual content; do not invent metrics. Keep each bullet to one sentence. Return bullets prefixed with a dash."
    user = f"Job Description:\n{jd_text}\n\nBullets to polish (keep meaning & metrics, improve clarity):\n{role_block_text}\n\nRewrite bullets (max 3–4), one sentence each, plain text."
    return call_openrouter(
        [{"role":"system","content":system},{"role":"user","content":user}],
        model=model, temperature=temperature, max_tokens=400, api_key=api_key
    )

# --------------------------
# UI
# --------------------------
st.title("ATS Resume Optimizer (≤2 pages, 5–6 line Profile)")

st.markdown("""Upload your baseline **DOCX** resume and paste a **Job Description**.
The app will:
- Extract JD keywords and compute **baseline ATS score**
- Rewrite **PROFILE** to 5–6 lines
- Trim sections to keep the resume within **2 pages**
- Keep original fonts/styles as much as possible (in-place edits)
- Produce an **optimized DOCX** + **ATS report**""")

col1, col2 = st.columns([1,1])
with col1:
    resume_file = st.file_uploader("Upload baseline resume (.docx)", type=["docx"], accept_multiple_files=False, key="resume")
with col2:
    jd = st.text_area("Paste Job Description", height=280, placeholder="Paste the job description here...")

advanced = st.expander("Advanced Options")
with advanced:
    top_k = st.slider("Max keywords from JD", 20, 80, 44, step=2)
    bullets_per_role = st.slider("Bullets per role (cap)", 2, 5, 3, step=1)
    projects_keep = st.slider("Projects to keep", 2, 6, 4, step=1)
    awards_keep = st.slider("Awards to keep", 2, 6, 4, step=1)
    hard_char_cap = st.slider("Soft character cap (approx 2 pages)", 7000, 12000, 9500, step=500)

    st.subheader("LLM Polishing (Optional)")
    use_llm = st.checkbox("Use OpenRouter to polish Profile & bullets", value=False, help="Requires OPENROUTER_API_KEY in Streamlit secrets")
    llm_model = st.selectbox("LLM Model (OpenRouter)", [
        "openrouter/auto",
        "openai/gpt-oss-20b:free",
        "deepseek/deepseek-r1-0528:free",
    ], index=0)
    llm_temp = st.slider("LLM temperature", 0.0, 1.0, 0.2, 0.1)

run = st.button("Optimize Resume")

if run:
    if not resume_file or not jd.strip():
        st.error("Please upload a DOCX resume and paste a job description.")
        st.stop()

    # Load resume
    doc = Document(resume_file)
    baseline_text = read_docx_text(doc)

    # Keywords & baseline score
    jd_keywords = extract_keywords(jd, top_k=int(top_k))
    base_score, base_matched = score_text(baseline_text, jd_keywords)

    # Build compact profile (or LLM rewrite)
    profile_text = build_compact_summary(jd_keywords)
    if 'OPENROUTER_API_KEY' in st.secrets and use_llm:
        try:
            profile_text = rewrite_summary_with_llm(profile_text, jd, st.secrets['OPENROUTER_API_KEY'], llm_model, llm_temp)
        except Exception as e:
            st.warning(f"LLM summary rewrite skipped: {e}")

    # Apply edits
    replace_profile_summary(doc, profile_text)
    cap_bullets_each_role(doc, n_keep=int(bullets_per_role))
    trim_projects(doc, keep=int(projects_keep))
    trim_awards(doc, keep=int(awards_keep))
    merge_profdev_interests(doc)
    enforce_two_pages_soft(doc, hard_char_cap=int(hard_char_cap))

    # Re-score
    final_text = read_docx_text(doc)
    final_score, final_matched = score_text(final_text, jd_keywords)

    # Optional: LLM bullet polishing (for copy-paste display)
    if 'OPENROUTER_API_KEY' in st.secrets and use_llm:
        try:
            full_txt = final_text
            start_kw = "EMPLOYMENT EXPERIENCE"
            end_markers = ["PERSONAL PROJECTS","PROJECTS","PROFESSIONAL DEVELOPMENT","CERTIFICATIONS","ACCOMPLISHMENTS","AWARDS","INTERESTS"]
            start_pos = full_txt.find(start_kw)
            if start_pos != -1:
                end_pos_candidates = [full_txt.find(m, start_pos) for m in end_markers]
                end_pos_candidates = [p for p in end_pos_candidates if p != -1]
                end_pos = min(end_pos_candidates) if end_pos_candidates else len(full_txt)
                role_block = full_txt[start_pos:end_pos]
                polished = polish_bullets_with_llm(role_block, jd, st.secrets['OPENROUTER_API_KEY'], llm_model, llm_temp)
                st.divider()
                st.subheader("Polished Bullets (copy-paste into your DOCX if you like)")
                st.code(polished)
        except Exception as e:
            st.warning(f"LLM bullet polishing skipped: {e}")

    # Assemble downloads
    out_buf = BytesIO()
    doc.save(out_buf)
    out_buf.seek(0)

    # Report
    missing = [k for k in jd_keywords if k not in final_matched]
    report = f"""# ATS Report
**Generated:** {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

## Summary
- **Baseline ATS score:** {base_score}%
- **Optimized ATS score:** {final_score}%

## JD Keywords (Top {len(jd_keywords)})
{", ".join(jd_keywords)}

## Matched after optimization ({len(final_matched)}):
{", ".join(sorted(final_matched))}

## Still Missing:
{", ".join(missing) if missing else "—"} 

## Changes Applied
- Rewrote **PROFILE** to a compact 5–6-line summary
- Capped bullets per role to **{bullets_per_role}**
- Kept only **{projects_keep}** projects and **{awards_keep}** awards
- Merged *Professional Development* and *Interests*
- Light-length control to target ≤2 pages
"""

    st.success("Optimization complete. Download your files below.")
    st.download_button("⬇️ Download Optimized Resume (DOCX)", out_buf, file_name="optimized_resume.docx")
    st.download_button("⬇️ Download ATS Report (Markdown)", report.encode("utf-8"), file_name="ATS_report.md")
    st.markdown(f"**Baseline ATS:** {base_score}%  |  **Optimized ATS:** {final_score}%")
    st.caption("Note: Scores are based on keyword presence and fuzzy matches; this is an approximation of ATS scanning.")
else:
    st.info("Upload your resume, paste the JD, then click **Optimize Resume**.")
